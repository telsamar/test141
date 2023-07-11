from telethon.sync import TelegramClient
from datetime import datetime, timedelta, timezone
from telethon.tl.types import MessageMediaDocument
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import os
from config import api_id, api_hash

filename = 'telegram_messages.docx'
screenshot_path = "screen"

chrome_options = Options()
chrome_options.add_argument("--headless")

webdriver_service = Service(ChromeDriverManager().install())

with open('dictionary.txt', 'r', encoding='utf-8') as f:
    words = [word.strip().lower() for word in f.readlines()]

try:
    doc = Document(filename)
except:
    doc = Document()

with TelegramClient('my_account3', api_id, api_hash) as client:
    time_limit = 15
    nomer = 1

    with open('channels.txt', 'r') as channels_file:
        channels = [channel.strip() for channel in channels_file.readlines()]
        
    for channel_name in channels:
        for message in client.iter_messages(channel_name):
            time_difference = (datetime.now(timezone.utc) - message.date)
            time_difference_seconds = time_difference.total_seconds()
            if time_difference_seconds < time_limit*3600:
                if any(word in message.text.lower() for word in words):
                    post_link = f"https://t.me/{channel_name}/{message.id}"

                    doc.add_paragraph(f'Nomer: {nomer}')
                    doc.add_paragraph(f'Message: {message.text}')
                    doc.add_paragraph(f'Link: {post_link}')

                    if isinstance(message.media, MessageMediaDocument):
                        if any(attribute for attribute in message.media.document.attributes if attribute.supports_streaming):
                            doc.add_paragraph('Есть видео')
                    else:
                        doc.add_paragraph('Нет!!!!!!!!!!')

                    driver = webdriver.Chrome(service=webdriver_service, options=chrome_options)
                    driver.get(post_link)
                    driver.set_window_size(1500, 3000)

                    element = driver.find_element(By.CLASS_NAME, "tgme_page_widget")
                    ActionChains(driver).move_to_element(element).perform()

                    screenshot_filename = screenshot_path + f'/screenshot_{nomer}.png'
                    element.screenshot(screenshot_filename)
                    driver.quit()

                    doc.add_picture(screenshot_filename, width=Inches(5))

                    os.remove(screenshot_filename)

                    nomer += 1
            else:
                break

doc.save(filename)
