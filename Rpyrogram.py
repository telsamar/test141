import asyncio
import logging
import os
import re
import sys
from datetime import datetime, timezone

from config import api_id, api_hash
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches
from pyrogram import Client
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
message_count = 0

# Определение временного лимита из аргументов командной строки или установка значения по умолчанию
if len(sys.argv) > 1:
    time_limit = int(sys.argv[1])
else:
    time_limit = 24

# Функция для чтения регулярных выражений из файла
def read_patterns(file_name):
    with open(file_name) as f:
        return [re.compile(pattern.strip()) for pattern in f.readlines()]

# Функция для чтения списка каналов из файла
def read_channels(file_name):
    with open(file_name, 'r') as f:
        return [line.strip().lower() for line in f.readlines()]

# Попытка чтения регулярных выражений и списка каналов
try:
    patterns = read_patterns('regex_patterns.txt')
    channels = read_channels('channels.txt')
    print(patterns)
except Exception as e:
    logger.error(f'Error reading files: {str(e)}')
    sys.exit(1)

# Определение пути для сохранения документа и скриншотов
filename = f'telegram_messages_{message_count}.docx'
screenshot_path = "screen"

# Настройка параметров веб-драйвера
chrome_options = Options()
chrome_options.add_argument("--headless")
webdriver_service = Service(ChromeDriverManager().install())

# Попытка открытия существующего документа или создания нового
try:
    doc = Document(filename)
    doc.clear()
except Exception:
    doc = Document()

# Основная  функция для работы с Telegram API и сохранения сообщений в документ Word
async def main():
    # Инициализация клиента Telegram с помощью Pyrogram
    app = Client("my_account2", api_id=api_id, api_hash=api_hash)
    await app.start()
    logger.info("Программа запущена...")

    # Добавление временного лимита в Word документ
    doc.add_paragraph(f'time_limit: {time_limit}')
    doc.save(filename)
    global message_count

    # Перебор списка каналов для обработки сообщений
    for channel_name in channels:
        async for message in app.get_chat_history(channel_name):
            # Вычисление разницы времени между текущим моментом и временем сообщения
            time_difference = datetime.now(timezone.utc) - message.date.astimezone(timezone.utc)
            time_difference_seconds = time_difference.total_seconds()
            # Проверка соответствия разницы времени установленному лимиту
            if time_difference_seconds < time_limit * 3600:
                post_link = f"https://t.me/{channel_name}/{message.id}"

                # Проверка на наличие совпадений с регулярными выражениями в тексте или подписи сообщения
                text_matches = message.text and any(pattern.search(message.text) for pattern in patterns)
                caption_matches = message.caption and any(pattern.search(message.caption) for pattern in patterns)

                if text_matches or caption_matches:
                    table = doc.add_table(rows=1, cols=2)
                    column_names = ["Номер", "Ссылка"]

                    for i, name in enumerate(column_names):
                        cell = table.cell(0, i)
                        cell.text = name

                    table.columns[0].width = Inches(1.0/3)
                    border_str = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single"/><w:left w:val="single"/><w:bottom w:val="single"/><w:right w:val="single"/></w:tcBorders>'

                    for cell in table.rows[0].cells:
                        border_elm = parse_xml(border_str)
                        cell._element.get_or_add_tcPr().append(border_elm)

                    message_count += 1
                    doc.add_paragraph(f'Nomer: {message_count}')
                    doc.add_paragraph(f'Text: {message.text}')
                    doc.add_paragraph(f'Caption: {message.caption}')
                    doc.add_paragraph(f'Link: {post_link}')

                    cells = table.add_row().cells
                    cells[0].text = str(message_count)
                    cells[1].text = post_link

                    for cell in cells:
                        border_elm = parse_xml(border_str)
                        cell._element.get_or_add_tcPr().append(border_elm)

                    logger.info(f'Nomer: {message_count}')
                    logger.info(f"Новое сообщение в канале: {message.chat.title}")
                    logger.info(f"Текст сообщения: {message.text}")
                    logger.info(f'Caption: {message.caption}')
                    logger.info(f'Link: {post_link}')

                    if message.photo:
                        logger.info("Есть фото")
                        doc.add_paragraph('Есть фото')
                    else:
                        logger.info("Нет фото")
                        doc.add_paragraph('Нет фото!!!!!!!!!!!!!!')

                    if message.video:
                        logger.info("Есть видео")
                        doc.add_paragraph('Есть видео')
                    else:
                        logger.info("Нет видео")
                        doc.add_paragraph('Нет видео!!!!!!!!!!!!!!')

                    if message.document:
                        logger.info("Есть прикрепленный документ")
                        doc.add_paragraph('Есть прикрепленный документ')
                    else:
                        logger.info("Нет прикрепленного документа")
                        doc.add_paragraph('Нет прикрепленного документа!!!!!!!!!!!!!!')

                    driver = webdriver.Chrome(service=webdriver_service, options=chrome_options)
                    driver.get(post_link)
                    driver.set_window_size(1500, 3000)
                    element = driver.find_element(By.CLASS_NAME, "tgme_page_widget")
                    ActionChains(driver).move_to_element(element).perform()

                    screenshot_filename = screenshot_path + f'/screenshot_{message_count}.png'
                    element.screenshot(screenshot_filename)
                    driver.quit()
                    doc.add_picture(screenshot_filename, width=Inches(5))
                    doc.save(filename)
                    os.remove(screenshot_filename)
            else:
                break

    print("Ожидание завершения - 3 секунды")
    await asyncio.sleep(3)
    await app.stop()


loop = asyncio.new_event_loop()
asyncio.set_event_loop(loop)

try:
    loop.run_until_complete(main())
finally:
    pending = asyncio.all_tasks(loop=loop)
    for task in pending:
        task.cancel()
    loop.run_until_complete(asyncio.gather(*pending, return_exceptions=True))
    loop.close()
